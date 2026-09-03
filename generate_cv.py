import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_resume():
    doc = docx.Document()

    # Set tight single-page margins
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Color palette
    PRIMARY_COLOR = RGBColor(0x1B, 0x36, 0x5D)  # Dark Navy #1B365D
    HEADING_COLOR = RGBColor(0x1B, 0x36, 0x5D)  # #1B365D
    TEXT_COLOR = RGBColor(0x11, 0x18, 0x27)     # Almost black #111827
    LINK_COLOR = RGBColor(0x00, 0x56, 0xB3)     # Classic Link Blue #0056B3
    GRAY_COLOR = RGBColor(0x4B, 0x55, 0x63)     # Gray #4B5563

    def add_p_bottom_border(p, color_hex="1B365D", sz="6"):
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{sz}" w:space="2" w:color="{color_hex}"/></w:pBdr>')
        pPr.append(pBdr)

    def add_hyperlink(paragraph, url, text, color=LINK_COLOR, underline=True):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="{r_id}" w:history="1"/>')
        new_run = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="{color[0]:02X}{color[1]:02X}{color[2]:02X}"/>{("<w:u w:val=\"single\"/>" if underline else "")}</w:rPr><w:t>{text}</w:t></w:r>')
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    # 1. NAME HEADER
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_name = p_name.add_run("HARSH GOYAL")
    run_name.font.name = "Calibri"
    run_name.font.size = Pt(19)
    run_name.font.bold = True
    run_name.font.color.rgb = PRIMARY_COLOR

    # 2. CONTACT INFO (2-Column Grid / Table)
    table_contact = doc.add_table(rows=2, cols=2)
    table_contact.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_contact.autofit = False

    for row in table_contact.rows:
        row.cells[0].width = Inches(4.3)
        row.cells[1].width = Inches(3.2)
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
            tcPr.append(tcBorders)

    # Row 0: LinkedIn (Left) | Email (Right)
    p_c0 = table_contact.rows[0].cells[0].paragraphs[0]
    p_c0.paragraph_format.space_before = Pt(0)
    p_c0.paragraph_format.space_after = Pt(1)
    r = p_c0.add_run("LinkedIn:  ")
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = TEXT_COLOR
    add_hyperlink(p_c0, "https://www.linkedin.com/in/harshgoyalcodes/", "https://www.linkedin.com/in/harshgoyalcodes/")

    p_c1 = table_contact.rows[0].cells[1].paragraphs[0]
    p_c1.paragraph_format.space_before = Pt(0)
    p_c1.paragraph_format.space_after = Pt(1)
    p_c1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_c1.add_run("Email: ")
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = TEXT_COLOR
    add_hyperlink(p_c1, "mailto:harshgoyal190720@gmail.com", "harshgoyal190720@gmail.com")

    # Row 1: GitHub (Left) | Mobile (Right)
    p_c2 = table_contact.rows[1].cells[0].paragraphs[0]
    p_c2.paragraph_format.space_before = Pt(0)
    p_c2.paragraph_format.space_after = Pt(3)
    r = p_c2.add_run("GitHub:    ")
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = TEXT_COLOR
    add_hyperlink(p_c2, "https://github.com/harshgoyal190720", "https://github.com/harshgoyal190720")

    p_c3 = table_contact.rows[1].cells[1].paragraphs[0]
    p_c3.paragraph_format.space_before = Pt(0)
    p_c3.paragraph_format.space_after = Pt(3)
    p_c3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_c3.add_run("Mobile: ")
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = TEXT_COLOR
    r_val = p_c3.add_run("+91 94642188078")
    r_val.font.name = "Calibri"
    r_val.font.size = Pt(9.5)
    r_val.font.color.rgb = TEXT_COLOR

    def add_section_title(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title_text)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = HEADING_COLOR
        add_p_bottom_border(p, "1B365D", "8")
        return p

    def add_bullet(p, bold_prefix, text_content):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.02
        p.paragraph_format.left_indent = Inches(0.2)
        
        r_bullet = p.add_run("•  ")
        r_bullet.font.name = "Calibri"
        r_bullet.font.size = Pt(9.2)
        r_bullet.font.color.rgb = TEXT_COLOR
        
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = "Calibri"
            r_bold.font.size = Pt(9.2)
            r_bold.font.bold = True
            r_bold.font.color.rgb = TEXT_COLOR

        r_text = p.add_run(text_content)
        r_text.font.name = "Calibri"
        r_text.font.size = Pt(9.2)
        r_text.font.color.rgb = TEXT_COLOR

    # =========================================================================
    # 3. SKILLS SECTION
    # =========================================================================
    add_section_title("SKILLS")
    skills = [
        ("Languages:  ", "C, C++, Python, JavaScript, HTML, CSS, SQL"),
        ("Technologies:  ", "React.js, Tailwind CSS, Google Authentication, Node.js fundamentals, Embedded C++, IoT"),
        ("Databases/Tools:  ", "MySQL, Git, GitHub, VS Code, Arduino IDE, Notion, Figma"),
        ("Soft Skills:  ", "Problem solving, Team collaboration, Technical mentorship, Content creation, Time management")
    ]
    for label, val in skills:
        p = doc.add_paragraph()
        add_bullet(p, label, val)

    # =========================================================================
    # 4. PROJECTS SECTION
    # =========================================================================
    add_section_title("PROJECTS")

    def add_item_header(title, gh_link_text, gh_url, date_str):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.rows[0].cells[0].width = Inches(5.7)
        table.rows[0].cells[1].width = Inches(1.8)
        
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
            tcPr.append(tcBorders)
            
        p_left = table.rows[0].cells[0].paragraphs[0]
        p_left.paragraph_format.space_before = Pt(2.5)
        p_left.paragraph_format.space_after = Pt(1)
        r_t = p_left.add_run(title)
        r_t.font.name = "Calibri"
        r_t.font.size = Pt(9.8)
        r_t.font.bold = True
        r_t.font.color.rgb = PRIMARY_COLOR
        
        if gh_link_text and gh_url:
            r_pipe = p_left.add_run(" | ")
            r_pipe.font.name = "Calibri"
            r_pipe.font.size = Pt(9.8)
            r_pipe.font.bold = True
            r_pipe.font.color.rgb = TEXT_COLOR
            add_hyperlink(p_left, gh_url, gh_link_text)
            
        p_right = table.rows[0].cells[1].paragraphs[0]
        p_right.paragraph_format.space_before = Pt(2.5)
        p_right.paragraph_format.space_after = Pt(1)
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_d = p_right.add_run(date_str)
        r_d.font.name = "Calibri"
        r_d.font.size = Pt(9.2)
        r_d.font.color.rgb = GRAY_COLOR

    # --- Project 1: PlaceIQ ---
    add_item_header("PlaceIQ (AI-Assisted Career Readiness Platform)", "GitHub", "https://github.com/harshgoyal190720/PlaceIQ", "May 2026 - Jun 2026")
    bullets_p1 = [
        ("Built an AI-assisted career readiness platform ", "designed to evaluate student profiles against target company skill benchmarks."),
        ("Engineered automated resume parsing ", "and skill-gap analysis pipelines identifying core technical deficiencies prior to job applications."),
        ("Structured tailored 6-step interview pathways ", "integrating company-specific profile assessments and interactive mock evaluations."),
        ("Tech Stack: ", "Python, NLP, React.js, JavaScript, Tailwind CSS, REST APIs, Git")
    ]
    for b_prefix, b_text in bullets_p1:
        p = doc.add_paragraph()
        add_bullet(p, b_prefix, b_text)

    # --- Project 2: TGPA Calculator ---
    add_item_header("TGPA Calculator (Semester Academic Utility)", "GitHub", "https://github.com/harshgoyal190720/coxmic-jassi-academic-calculator", "Dec 2025 - Jan 2026")
    bullets_p2 = [
        ("Developed a responsive web application ", "enabling students to accurately calculate semester TGPA and projected CGPA metrics."),
        ("Integrated Google OAuth 2.0 authentication ", "allowing secure user login and persistent cloud/local grade record synchronization."),
        ("Designed dynamic credit-grade calculation algorithms ", "supporting various university grading scales and instant GPA projections."),
        ("Tech Stack: ", "React.js, JavaScript, Tailwind CSS, Google OAuth 2.0, HTML5, LocalStorage")
    ]
    for b_prefix, b_text in bullets_p2:
        p = doc.add_paragraph()
        add_bullet(p, b_prefix, b_text)

    # --- Project 3: Arduino Surveillance System ---
    add_item_header("Arduino Based Surveillance & Monitoring System", "GitHub", "https://github.com/harshgoyal190720/arduinosurveillance", "Nov 2025 - Dec 2025")
    bullets_p3 = [
        ("Engineered a compact embedded security system ", "combining ultrasonic proximity detection and environmental monitoring with Arduino UNO."),
        ("Integrated HC-SR04 SONAR (2–400cm range) ", "and DHT11 temperature/humidity sensors with active buzzer alarm on Digital Pin 8 polling every 500ms."),
        ("Implemented I2C 16x2 LCD display readouts ", "at address 0x27, achieving 99% real-time threshold detection accuracy."),
        ("Tech Stack: ", "Arduino UNO, Embedded C++, HC-SR04 Ultrasonic Sensor, DHT11, I2C LCD, Active Buzzer")
    ]
    for b_prefix, b_text in bullets_p3:
        p = doc.add_paragraph()
        add_bullet(p, b_prefix, b_text)

    # =========================================================================
    # 5. CERTIFICATES SECTION
    # =========================================================================
    add_section_title("CERTIFICATES")
    certs = [
        ("Programming Fundamentals using Python - Part 2 | ", "Infosys Springboard", "Jul 2026"),
        ("Design Thinking and Innovation | ", "IIT Bombay & Coursera", "Jul 2026"),
        ("Computer Programming in C (150 Hours Course Duration) | ", "NeoColab • iamneo & LPU", "May 2026"),
        ("ESL002: Intermediate English as a Second Language | ", "Saylor Academy", "Jan 2026")
    ]
    for c_title, c_issuer, c_date in certs:
        table_c = doc.add_table(rows=1, cols=2)
        table_c.alignment = WD_TABLE_ALIGNMENT.LEFT
        table_c.autofit = False
        table_c.rows[0].cells[0].width = Inches(6.1)
        table_c.rows[0].cells[1].width = Inches(1.4)
        for cell in table_c.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
            tcPr.append(tcBorders)
            
        p_c_left = table_c.rows[0].cells[0].paragraphs[0]
        p_c_left.paragraph_format.space_before = Pt(0)
        p_c_left.paragraph_format.space_after = Pt(1)
        r_dot = p_c_left.add_run("•  ")
        r_dot.font.name = "Calibri"
        r_dot.font.size = Pt(9.2)
        r_ct = p_c_left.add_run(c_title)
        r_ct.font.name = "Calibri"
        r_ct.font.size = Pt(9.2)
        r_ct.font.bold = True
        r_ct.font.color.rgb = TEXT_COLOR
        r_ci = p_c_left.add_run(c_issuer)
        r_ci.font.name = "Calibri"
        r_ci.font.size = Pt(9.2)
        r_ci.font.color.rgb = PRIMARY_COLOR
        
        p_c_right = table_c.rows[0].cells[1].paragraphs[0]
        p_c_right.paragraph_format.space_before = Pt(0)
        p_c_right.paragraph_format.space_after = Pt(1)
        p_c_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_cd = p_c_right.add_run(c_date)
        r_cd.font.name = "Calibri"
        r_cd.font.size = Pt(9.2)
        r_cd.font.color.rgb = GRAY_COLOR

    # =========================================================================
    # 6. ACHIEVEMENTS SECTION
    # =========================================================================
    add_section_title("ACHIEVEMENTS")
    achievements = [
        ("Academic Distinction: ", "Maintained a top-tier CGPA of 9.31 in B.Tech Computer Science & Engineering at LPU."),
        ("Technical Content Creator: ", "Founded Harsh Codes YouTube channel; authored Python 21-Day Bootcamp & CSE Placement Guides."),
        ("Rigorous Problem Solving: ", "Completed 150+ hours of structured programming in C, C++, and Python."),
        ("Hands-on Engineering: ", "Architected 5+ full-stack, AI, and IoT hardware projects independently.")
    ]
    for a_prefix, a_text in achievements:
        p = doc.add_paragraph()
        add_bullet(p, a_prefix, a_text)

    # =========================================================================
    # 7. EDUCATION SECTION
    # =========================================================================
    add_section_title("EDUCATION")

    def add_edu_item(inst_name, location, degree_info, date_range):
        table_e = doc.add_table(rows=2, cols=2)
        table_e.alignment = WD_TABLE_ALIGNMENT.LEFT
        table_e.autofit = False
        table_e.rows[0].cells[0].width = Inches(5.2)
        table_e.rows[0].cells[1].width = Inches(2.3)
        table_e.rows[1].cells[0].width = Inches(5.2)
        table_e.rows[1].cells[1].width = Inches(2.3)
        
        for row in table_e.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
                tcPr.append(tcBorders)
                
        p_e0 = table_e.rows[0].cells[0].paragraphs[0]
        p_e0.paragraph_format.space_before = Pt(1.5)
        p_e0.paragraph_format.space_after = Pt(0.5)
        r_bullet = p_e0.add_run("•  ")
        r_bullet.font.name = "Calibri"
        r_bullet.font.size = Pt(9.2)
        r_in = p_e0.add_run(inst_name)
        r_in.font.name = "Calibri"
        r_in.font.size = Pt(9.2)
        r_in.font.bold = True
        r_in.font.color.rgb = PRIMARY_COLOR
        
        p_e1 = table_e.rows[0].cells[1].paragraphs[0]
        p_e1.paragraph_format.space_before = Pt(1.5)
        p_e1.paragraph_format.space_after = Pt(0.5)
        p_e1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_loc = p_e1.add_run(location)
        r_loc.font.name = "Calibri"
        r_loc.font.size = Pt(9.2)
        r_loc.font.color.rgb = GRAY_COLOR
        
        p_e2 = table_e.rows[1].cells[0].paragraphs[0]
        p_e2.paragraph_format.space_before = Pt(0)
        p_e2.paragraph_format.space_after = Pt(2)
        p_e2.paragraph_format.left_indent = Inches(0.2)
        r_deg = p_e2.add_run(degree_info)
        r_deg.font.name = "Calibri"
        r_deg.font.size = Pt(9.2)
        r_deg.font.color.rgb = TEXT_COLOR
        
        p_e3 = table_e.rows[1].cells[1].paragraphs[0]
        p_e3.paragraph_format.space_before = Pt(0)
        p_e3.paragraph_format.space_after = Pt(2)
        p_e3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_dt = p_e3.add_run(date_range)
        r_dt.font.name = "Calibri"
        r_dt.font.size = Pt(9.2)
        r_dt.font.color.rgb = GRAY_COLOR

    add_edu_item(
        "Lovely Professional University",
        "Phagwara, Punjab",
        "Bachelor of Technology - Computer Science and Engineering; CGPA: 9.31",
        "Aug 2025 - Present"
    )
    add_edu_item(
        "Bhai Mastan Singh Public School",
        "Muktsar, Punjab",
        "Higher Secondary Education (12th); Percentage: 75.8",
        "Mar 2024 - Mar 2025"
    )
    add_edu_item(
        "D.A.V. Public School",
        "Muktsar, Punjab",
        "Secondary Education (10th); Percentage: 90.6",
        "Mar 2022 - Mar 2023"
    )

    out_path = "Harsh_Goyal_Resume.docx"
    doc.save(out_path)
    doc.save("assets/Harsh_Goyal_Resume.docx")
    print(f"Resume generated successfully at: {os.path.abspath(out_path)}")

if __name__ == "__main__":
    create_resume()
